from django.http import HttpResponse
from docx import Document
from docx.oxml.shape import CT_Inline
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from io import BytesIO
from docx.enum.text import WD_LINE_SPACING
import program.word.jalali
from django.conf import settings


def styling(c, t, f):
    p1 = c.add_paragraph('')
    run = p1.add_run(t)
    run.bold = True
    p1_format = p1.paragraph_format
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(f)
    p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p1


def styling12(c, t, f):
    p1 = c.add_paragraph('')
    run = p1.add_run(t)
    run.bold = True
    p1_format = p1.paragraph_format
    font = run.font
    font.name = 'B Nazanin'
    font.size = Pt(f)
    p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p1


def styling1(c, t, f):
    p1 = c.add_paragraph('')
    run = p1.add_run(t)
    p1_format = p1.paragraph_format
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(f)
    p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p1


def add_picture_to_run(run, picture_file, width=None, height=None):
    """
    Add a picture at the end of a run.
    """
    rId, image = run.part.get_or_add_image(picture_file)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = run.part.next_id, image.filename
    inline = CT_Inline.new_pic_inline(shape_id, rId, filename, cx, cy)
    run._r.add_drawing(inline)







def registrations_to_print(registerations):
    docx_title = "print.docx"
    document = Document()
    section = document.sections[-1]
    section.right_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.top_margin = Cm(0.5)
    section.down_margin = Cm(0.5)
    section.page_width = Cm(10.5)
    section.page_height = Cm(14.85)

    for item in registerations:
        p = document.add_paragraph('نام و نام خانوادگی :')
        # p.bold = True
        # .bold = True
        run = p.add_run(str(item.profile.user.first_name))
        font = run.font
        run.bold = True
        font.name = 'B Nazanin'
        p.add_run(str(item.profile.user.last_name)).bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = document.add_paragraph('کد ملی')
        p.add_run(str(item.profile.user.username ) + ':').bold = True
        p.add_run(' شماره شناسنامه')
        p.add_run('  ')
        p.add_run(str(item.profile.shenasname) + ':').bold = True
        p = document.add_paragraph('تلفن همراه')
        p.add_run(str(item.profile.cellPhone) + ':').bold = True
        p.add_run(' تلفن ضروری')
        p.add_run('  ')
        p.add_run(str(item.profile.emergencyPhone) + ':').bold = True
        p_format = p.paragraph_format
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        a = ''
        if item.profile.people_type == 'sharif graduated':
            a = 'فارغ التحصیل شریف'
        if item.profile.people_type == 'sharif student':
            a = 'دانشجو شریف'
        if item.profile.people_type == 'sharif graduated talabe':
            a = 'فاغ التحصیل شریف و طلبه فعلی'
        if item.profile.people_type == 'sharif graduated student not sharif':
            a = 'فارغ التحصیل شریف و دانشجو سایر'
        if item.profile.people_type == 'not sharif student':
            a = 'دانشجو سایر'
        if item.profile.people_type == 'not sharif graduated':
            a = 'فارغ التحصیل سایر'
        if item.profile.people_type == 'talabe':
            a = 'طلبه'
        if item.profile.people_type == 'sharif master':
            a = 'استاد شریف'
        if item.profile.people_type == 'sharif employed':
            a = 'کارمند شریف'
        if item.profile.people_type == 'other':
            a = 'سایر'
        p = document.add_paragraph('')
        p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.add_run(str(a)).bold = True
        if item.coupling:
            a = 'متاهلی'
        else:
            a = 'مجردی'
        p.add_run(' ' + str(a) + ' ').bold = True
        if item.additionalOption:
            b1 = 'هوایی'
        else:
            b1 = 'زمینی'
        p.add_run(b1).bold = True
        p.add_run(' شماره دانشجویی')
        p.add_run('  ')
        p.add_run(str(item.profile.studentNumber) + ':').bold = True
        # p.add_run(str(item.profile.user.last_name)).bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p = document.add_paragraph('')

        # if item.profile. Mugshot:
        #     a = 'عکس دارد'
        run = p.add_run('عکس دارد')
        picture = open(settings.BASE_DIR+'/program/word/3.JPG', 'rb')
        add_picture_to_run(run, picture, width=Cm(0.5), height=Cm(0.5))
        p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.add_run('  ')
        run = p.add_run('کپی کارت ملی')
        add_picture_to_run(run, picture, width=Cm(0.5), height=Cm(0.5))
        p.add_run('  ')
        p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        run = p.add_run('تعهد نامه')
        add_picture_to_run(run, picture, width=Cm(0.5), height=Cm(0.5))
        p_format = p.paragraph_format
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = document.add_paragraph('وضعیت نظام وظیفه :')
        b = ''
        C = ''
        D = ''
        if item.profile.conscription == 'went':
            b = 'دارای کارت پایان خدمت'
            C = 'کپی کارت پایان خدمت دارد'
        if item.profile.conscription == 'exempt':
            b = 'معافیت دایم '
            C = 'کپی کارت معافیت'
        if item.profile.conscription == 'educational exempt':
            b = 'معافیت تحصیلی'
            C = 'فرم اشتغال به تحصیل'
            D = ' تعهد خروج از کشور '
        p.add_run(b).bold = True

        p_format = p.paragraph_format
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = document.add_paragraph(C)
        run = p.add_run('  ')
        add_picture_to_run(run, picture, width=Cm(0.5), height=Cm(0.5))
        font = run.font
        font.name = 'B Nazanin'
        p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        run = p.add_run(str(D))
        if D == ' تعهد خروج از کشور ':
            add_picture_to_run(run, picture, width=Cm(0.5), height=Cm(0.5))
        # p.add_run(str(item.profile.user.last_name))

        p_format = p.paragraph_format
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = document.add_paragraph('وضیعت گذرنامه' + ':')
        b = ''
        C = ''
        if item.profile.passport_choices == 'not have':
            b = 'ندارم'
        if item.profile.passport_choices == 'have':
            b = 'دارم'
            C = 'تصویر صفحه اول گذر نامه'
        if item.profile.passport_choices == 'have 7':
            b = 'دارم ولی کمتر از 7 ماه انقضا دارد '
        p.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.add_run(str(b)).blod = True
        p.add_run('   ')
        run = p.add_run(str(C))
        if C == 'تصویر صفحه اول گذر نامه':
            add_picture_to_run(run, picture, width=Cm(0.5), height=Cm(0.5))
        p_format = p.paragraph_format
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = document.add_paragraph('وضیعت پرداخت' + ':')

        run = p.add_run('پرداخت اول')
        add_picture_to_run(run, picture, width=Cm(0.5), height=Cm(0.5))
        p.add_run('    ')
        run = p.add_run('پرداخت دوم')
        add_picture_to_run(run, picture, width=Cm(0.5), height=Cm(0.5))
        # p.add_run(str(item.profile.user.last_name))

        p_format = p.paragraph_format
        p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_page_break()
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response


def registrations_to_manifest(registerations):
    docx_title = "manifest.docx"
    document = Document()
    section = document.sections[-1]
    section.right_margin = Cm(0.5)
    section.left_margin = Cm(0.5)
    section.top_margin = Cm(0.5)
    section.down_margin = Cm(0.5)

    table = document.add_table(1, 3)
    hdr_cells = table.rows[0].cells
    p1 = hdr_cells[1].add_paragraph('')

    run = p1.add_run('')
    picture = open(settings.BASE_DIR+'/program/word/1.jpg', 'rb')
    add_picture_to_run(run, picture, width=Cm(5.5), height=Cm(4))
    run.bold = True
    p1_format = p1.paragraph_format
    p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p1 = hdr_cells[2].add_paragraph('')
    run = p1.add_run('')
    picture = open(settings.BASE_DIR+'/program/word/2.jpg', 'rb')
    add_picture_to_run(run, picture, width=Cm(5.5), height=Cm(4))
    run.bold = True
    p1_format = p1.paragraph_format
    p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #
    p0 = hdr_cells[0].add_paragraph('')
    p11 = hdr_cells[0].add_paragraph('')
    styling1(hdr_cells[0], 'Embassy of the Republic of', 14)
    styling(hdr_cells[0], 'IRAQ', 14)
    styling(hdr_cells[0], 'TEHRAN', 14)

    run = p1.add_run('')
    run.bold = True
    p1_format = p1.paragraph_format
    font = run.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p1 = document.add_paragraph('')
    run = p1.add_run('العدد:')
    p1_format = p1.paragraph_format
    p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p2 = document.add_paragraph('التاریخ:')

    # create table and style
    table = document.add_table(2, 12, 'Table Grid')
    # create row number
    r1 = table.rows[0].cells
    r2 = table.rows[1].cells

    # size of rows
    r1[0].width = Cm(3.5)
    r1[1].width = Cm(2)
    r1[2].width = Cm(3)
    r1[3].width = Cm(1.5)
    r1[4].width = Cm(1)
    r1[5].width = Cm(1)
    r1[6].width = Cm(4)
    r1[7].width = Cm(3)
    r1[8].width = Cm(3)
    r1[9].width = Cm(4)
    r1[10].width = Cm(2)
    r1[11].width = Cm(2)

    # font and text
    a = r1[0]

    styling(r1[1], 'اصلی', 12)
    c = r1[3]
    d = r1[4]
    e = r1[5]
    B = c.merge(d).merge(e)
    styling(B, 'الموالید', 12)
    f = r1[6]
    styling(r1[7], 'رقم الکتاب', 12)

    b = r2[0]
    A = a.merge(b)
    styling(A, 'الصورة', 12)
    styling(r2[1], 'مرافق', 12)
    D = r2[2].merge(r1[2])
    styling(D, 'میلادی', 12)
    styling(r2[3], 'سنة', 12)
    styling(r2[4], 'شهر', 12)
    styling(r2[5], 'یوم', 12)
    h = r2[6]
    C = f.merge(h)
    styling(C, 'رقم الجواز و محل و تاریخ الصدور', 12)
    styling(r2[7], 'تاریخة', 12)
    styling(r2[8].merge(r1[8]), 'الجنس', 12)
    styling(r2[9].merge(r1[9]), 'الاسم الکامل', 12)
    styling(r2[10].merge(r1[10]), 'ت', 12)
    styling(r2[11].merge(r1[11]), 'تسلسل فی القائمة', 12)

    i = 1
    for item in registerations:

        if ((i - 1) % 5 == 0) and (i != 1):
            document.add_page_break()
            section = document.sections[-1]
            section.right_margin = Cm(0.5)
            section.left_margin = Cm(0.5)
            section.top_margin = Cm(0.5)
            section.down_margin = Cm(0.5)

            table = document.add_table(1, 3)
            hdr_cells = table.rows[0].cells
            p1 = hdr_cells[1].add_paragraph('')

            run = p1.add_run('')
            picture = open('./program/word/1.jpg', 'rb')
            add_picture_to_run(run, picture, width=Cm(5.5), height=Cm(4))
            run.bold = True
            p1_format = p1.paragraph_format
            p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p1 = hdr_cells[2].add_paragraph('')
            run = p1.add_run('')
            picture = open('./program/word/2.jpg', 'rb')
            add_picture_to_run(run, picture, width=Cm(5.5), height=Cm(4))
            run.bold = True
            p1_format = p1.paragraph_format
            p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #
            p0 = hdr_cells[0].add_paragraph('')
            p11 = hdr_cells[0].add_paragraph('')
            styling1(hdr_cells[0], 'Embassy of the Republic of', 14)
            styling(hdr_cells[0], 'IRAQ', 14)
            styling(hdr_cells[0], 'TEHRAN', 14)

            run = p1.add_run('')
            run.bold = True
            p1_format = p1.paragraph_format
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
            p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p1 = document.add_paragraph('')
            run = p1.add_run('العدد:')
            p1_format = p1.paragraph_format
            p1_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p2 = document.add_paragraph('التاریخ:')

            # create table and style
            table = document.add_table(2, 12, 'Table Grid')
            # create row number
            r1 = table.rows[0].cells
            r2 = table.rows[1].cells

            # size of rows
            r1[0].width = Cm(3.5)
            r1[1].width = Cm(2)
            r1[2].width = Cm(1)
            r1[3].width = Cm(1)
            r1[4].width = Cm(1)
            r1[5].width = Cm(1)
            r1[6].width = Cm(10)
            r1[7].width = Cm(3)
            r1[8].width = Cm(3)
            r1[9].width = Cm(4)
            r1[10].width = Cm(2)
            r1[11].width = Cm(2)

            # font and text
            a = r1[0]

            styling(r1[1], 'اصلی', 12)
            c = r1[3]
            d = r1[4]
            e = r1[5]
            B = c.merge(d).merge(e)
            styling(B, 'الموالید', 12)
            f = r1[6]
            styling(r1[7], 'رقم الکتاب', 12)

            b = r2[0]
            A = a.merge(b)
            styling(A, 'الصورة', 12)
            styling(r2[1], 'مرافق', 12)
            D = r2[2].merge(r1[2])
            styling(D, 'میلادی', 12)
            styling(r2[3], 'سنة', 12)
            styling(r2[4], 'شهر', 12)
            styling(r2[5], 'یوم', 12)
            h = r2[6]
            C = f.merge(h)
            styling(C, 'رقم الجواز و محل و تاریخ الصدور', 12)
            styling(r2[7], 'تاریخة', 12)
            styling(r2[8].merge(r1[8]), 'الجنس', 12)
            styling(r2[9].merge(r1[9]), 'الاسم الکامل', 12)
            styling(r2[10].merge(r1[10]), 'ت', 12)
            styling(r2[11].merge(r1[11]), 'تسلسل فی القائمة', 12)
        if item.profile.gender == 1:
            j = 'الذکر'
        elif item.profile.gender == 0:
            j = 'اثنی'
        row_cells = table.add_row().cells
        row_cells[10].text = str(i)
        row_cells[8].text = str(j)
        i = i + 1
        styling(row_cells[9], str(item.profile.user.first_name), 12)
        styling(row_cells[9], str(item.profile.fatherName), 12)
        styling(row_cells[9], str(item.profile.user.last_name), 12)
        a1 = str(item.profile.passport_number)
        if len(a1) == 9:
            a1 = a1[1:8]
        styling(row_cells[6], a1, 12)
        styling(row_cells[6], str('ایران'), 12)
        styling(row_cells[6], str(item.profile.passport_dateofissue), 12)
        styling(row_cells[6], str(item.profile.passport_dateofexpiry), 12)
        d = str(item.profile.birthDay)
        m = str(item.profile.birthMonth)
        y = str(item.profile.birthYear)
        t = '13' + y + '-' + m + '-' + d
        if len(y) == 4:
            t = y + '-' + m + '-' + d
        j1 = program.word.jalali.Persian(t).gregorian_day
        u1 = program.word.jalali.Persian(t).gregorian_month
        t1 = program.word.jalali.Persian(t).gregorian_year
        styling(row_cells[5], str(j1), 12)
        styling(row_cells[4], str(u1), 12)
        styling(row_cells[3], str(t1), 12)
        p1 = row_cells[0].add_paragraph('')
        if item.profile.mugshot:
          picture = open('./' + str(item.profile.mugshot), 'rb')
          run = p1.add_run('')
          add_picture_to_run(run, picture, width=Cm(2), height=Cm(3))
        else:
            picture = open('./program/word/4.jpg', 'rb')
            run = p1.add_run('')
            add_picture_to_run(run, picture, width=Cm(2), height=Cm(3))
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response
